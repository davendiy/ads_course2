#!/usr/bin/env python3
# -*-encoding: utf-8-*-

# created: 28.11.18
# by David Zashkolny
# 2 course, comp math
# Taras Shevchenko National University of Kyiv
# email: davendiy@gmail.com


import docx



def statistic(filename):
    document = docx.Document(filename)
    tables = document.tables        # таблиці документа
    paragraph_amount = len(document.paragraphs)   # к-ть параграфів (параграфи в таблицях не враховуються)
    symb_amount = 0    # к-ть символів
    word_amount = 0    # к-ть слів
    for paragraph in document.paragraphs:     # проходим по параграфам
        for run in paragraph.runs:            # в кожному параграфі по потокам
            symb_amount += len(run.text().strip().replace(' ', ''))   # убираєм \n і видаляєм пробіли
            word_amount += len(run.text().split())
    for column in tables.columns:
        for cell in column:
            symb_amount += len(cell.text().strip().replace(' ', ''))

    return paragraph_amount, symb_amount, word_amount


if __name__ == '__main__':
    print("statistic of input.docx: ", statistic('input.docx'))
