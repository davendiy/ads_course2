#!/usr/bin/env python3
# -*-encoding: utf-8-*-

from docx import Document
import os

from dz_others.subject23_MS.merge.mergesource import *
from dz_others.subject23_MS.merge.sourceitem import *


class Merger:

    def __init__(self, template, param_files, leadparam, outfile=None):
        self.indoc = Document(template)
        if not leadparam in param_files:   # С‚СЂРµР±Р°, С‰РѕР± РїСЂРѕРІС–РґРЅРёР№ РїР°СЂР°РјРµС‚СЂ РѕР±РѕРІ'СЏР·РєРѕРІРѕ Р±СѓРІ
            raise KeyError
        # Р·РІ'СЏР·Р°С‚Рё Р· РґР¶РµСЂРµР»Р°РјРё РґР°РЅРёС…, РїС–РґРіРѕС‚СѓРІР°С‚РёСЃСЏ РґРѕ Р·Р»РёС‚С‚СЏ
        self.mergesrc = MergeSource(param_files, leadparam)
        if not outfile:
            base, ext = os.path.splitext(template)  # Р±СѓРґСѓС”РјРѕ С–Рј'СЏ С„Р°Р№Р»Сѓ-СЂРµР·СѓР»СЊС‚Р°С‚Сѓ
            outfile = base + '_merged.docx'
        self.outfile = outfile
        self.outdoc = Document()

    def merge(self):
        for record in self.mergesrc:
            self._process_template(record)
            self.outdoc.add_page_break()  # РґРѕРґР°С‚Рё СЂРѕР·СЂРёРІ СЃС‚РѕСЂС–РЅРєРё
        self.outdoc.save(self.outfile)  # Р·Р±РµСЂРµРіС‚Рё С„Р°Р№Р» СЂРµР·СѓР»СЊС‚Р°С‚Сѓ

    def _process_template(self, record):
        for paragraph in self.indoc.paragraphs:
            self._process_paragraph(paragraph, record)

    def _process_paragraph(self, paragraph, record):
        new_para = self.outdoc.add_paragraph(style=paragraph.style)
        para_format_copy(paragraph, new_para)
        for run in paragraph.runs:
            self._process_run(run, record, new_para)

    def _process_run(self, run, record, new_para):
        text = run.text
        while True:
            pos, field = self._get_field_pos(text, record)  # С‡Рё С” РјС–С‚РєР° РїРѕР»СЏ РґР°РЅРёС…
            if pos == -1: break
            if pos > 0:  # СЏРєС‰Рѕ РґРѕ РјС–С‚РєРё С” С‚РµРєСЃС‚, РґРѕРґР°С”РјРѕ Р№РѕРіРѕ РґРѕ СЂРµР·СѓР»СЊС‚Р°С‚Сѓ
                new_run = new_para.add_run(text[:pos], run.style)
                run_format_copy(run, new_run)
            text = text[
                   pos + len(field) + 2:]  # РІРёРєР»СЋС‡Р°С”РјРѕ РґРѕРґР°РЅРёР№ С‚РµРєСЃС‚ С‚Р° РјС–С‚РєСѓ РїРѕР»СЏ
            srcitem = record[field]  # РѕС‚СЂРёРјСѓС”РјРѕ РѕР±'С”РєС‚ Р· РґР°РЅРёРјРё SourceItem
            if srcitem.type == 'excel':
                # РґР»СЏ Excel СѓСЃРµ РїСЂРѕСЃС‚Рѕ
                new_run = new_para.add_run(str(srcitem.obj.value), run.style)
                run_format_copy(run, new_run)
            else:
                num_para = len(srcitem.obj.paragraphs)
                for i, t_para in enumerate(srcitem.obj.paragraphs):
                    for t_run in t_para.runs:
                        new_run = new_para.add_run(t_run.text, t_run.style)
                        run_format_copy(t_run, new_run)
                    if i < num_para - 1:  # СЏРєС‰Рѕ РїР°СЂР°РіСЂР°С„ РЅРµ РѕСЃС‚Р°РЅРЅС–Р№, РґРѕРґР°С”РјРѕ СЂРѕР·СЂРёРІ СЂСЏРґРєР°
                        new_run.add_break()
        new_run = new_para.add_run(text, run.style)
        run_format_copy(run, new_run)

    def _get_field_pos(self, text, record):
        pos = -1
        field = None
        for name in record:
            s = '{' + name + '}'  # С„РѕСЂРјСѓС”РјРѕ СЂСЏРґРѕРє РјС–С‚РєРё
            pos = text.find(s)
            if pos >= 0:
                field = name
                break
        return pos, field


def para_format_copy(p1, p2):
    pf1 = p1.paragraph_format  # РѕР±'С”РєС‚ ParagraphFormat
    pf2 = p2.paragraph_format
    pf2.alignment = pf1.alignment
    pf2.line_spacing = pf1.line_spacing
    pf2.first_line_indent = pf1.first_line_indent
    pf2.keep_together = pf1.keep_together
    pf2.keep_with_next = pf1.keep_with_next
    pf2.left_indent = pf1.left_indent
    pf2.line_spacing_rule = pf1.line_spacing_rule
    pf2.page_break_before = pf1.page_break_before
    pf2.right_indent = pf1.right_indent
    pf2.space_after = pf1.space_after
    pf2.space_before = pf1.space_before
    pf2.widow_control = pf1.widow_control


def run_format_copy(r1, r2):
    r2.bold = r1.bold
    r2.italic = r1.italic
    r2.underline = r1.underline
    rf1 = r1.font  # РѕР±'С”РєС‚ Font
    rf2 = r2.font
    rf2.all_caps = rf1.all_caps
    rf2.color.rgb = rf1.color.rgb  # РѕР±'С”РєС‚ ColorFormat
    rf2.name = rf1.name
    rf2.size = rf1.size
    rf2.small_caps = rf1.small_caps
    rf2.strike = rf1.strike
    rf2.subscript = rf1.subscript
    rf2.superscript = rf1.superscript


if __name__ == '__main__':
    import sys

    if len(sys.argv) == 1:  # СЏРєС‰Рѕ РЅРµРјР°С” РїР°СЂР°РјРµС‚СЂС–РІ, СЃС‚Р°РІРёРјРѕ С–Рј'СЏ Р·Р° СѓРіРѕРґРѕСЋ
        config = "config.txt"
    else:
        config = sys.argv[1]  # 1 РїР°СЂР°РјРµС‚СЂ
    conf = ConfigDict(config)
    params = conf.getconfig()
    # РѕРєСЂРµРјРѕ РІРёРґС–Р»СЏС”РјРѕ leadparam, template, outfile
    leadparam = params['LeadParam']
    del params['LeadParam']
    template = params['Template']
    del params['Template']
    outfile = params.get('OutFile')
    if outfile:
        del params['OutFile']

    merger = Merger(template, params, leadparam, outfile)
    merger.merge()
