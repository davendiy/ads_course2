#!/usr/bin/env python3
# -*-encoding: utf-8-*-

# by David Zashkolny
# 2 course, comp math
# Taras Shevchenko National University of Kyiv
# email: davendiy@gmail.com


import os
from docx import Document


def format_docx(directory, etalon_name):

    filelist = os.listdir(directory)
    standard = Document(os.path.join(directory, etalon_name))

    for filename in filelist:
        if not filename.endswith('.docx'):
            continue

        tmp_full_name = os.path.join(directory, filename)
        tmp_doc = Document(tmp_full_name)
        for standard_par, file_par in zip(standard.paragraphs, tmp_doc.paragraphs):
            file_par.paragraph_format = standard_par.paragraph_format
            file_par.style = standard_par.style

            for standard_par_run, file_par_run in zip(standard_par.runs, file_par.runs):
                file_par_run.style = standard_par_run.style

        tmp_doc.save(tmp_full_name)


if __name__ == '__main__':
    path = os.getcwd()
    path = os.path.join(path, 't23_1')
    format_docx(path, 'Standard.docx')
